import os
import logging
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from google import genai
import json
import openai
from openai import OpenAI
import requests
from pylatex import Document, Section, Subsection, Command
from pylatex.utils import italic, NoEscape
import tempfile
import subprocess
from jinja2 import Template

logging.basicConfig(level=logging.INFO)
log = logging.getLogger("resume-backend")

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
ATS_API_KEY = os.getenv("ATS_API_KEY")  # For ATS scoring API

if not GEMINI_API_KEY:
    raise Exception("❌ No Gemini API key found in .env!")

if not OPENAI_API_KEY:
    raise Exception("❌ No OpenAI API key found in .env!")

gemini_client = genai.Client(api_key=GEMINI_API_KEY)
openai_client = OpenAI(api_key=OPENAI_API_KEY)
log.info("Gemini and OpenAI clients initialized.")
log.info(f"Gemini model: {GEMINI_MODEL}")

app = Flask(__name__)
CORS(app)


def enhance_with_openai(text: str) -> str:
    prompt = f"""
Enhance the following resume text.
- Improve grammar
- Rewrite into strong ATS-optimized bullets
- Use measurable achievements
- Improve tone
- Output ONLY improved resume text.

Resume:
{text}
"""

    response = openai_client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1000
    )
    return response.choices[0].message.content.strip()

def enhance_with_gemini(text: str) -> str:
    prompt = f"""
Enhance the following resume text.
- Improve grammar
- Rewrite into strong ATS-optimized bullets
- Use measurable achievements
- Improve tone
- Output ONLY improved resume text.

Resume:
{text}
"""

    response = gemini_client.models.generate_content(
        model=GEMINI_MODEL,
        contents=prompt
    )
    return response.text


def ats_analysis(text: str):
    # Integrate real free ATS API, e.g., Jobscan's free tier (requires API key)
    if ATS_API_KEY:
        try:
            response = requests.post(
                "https://api.jobscan.co/v1/ats-score",  # Replace with actual free ATS API endpoint if different
                headers={"Authorization": f"Bearer {ATS_API_KEY}"},
                json={"resume": text}
            )
            if response.status_code == 200:
                data = response.json()
                return {
                    "overall_score": data.get("overall_score", 78),
                    "keyword_score": data.get("keyword_score", 71),
                    "skill_match_score": data.get("skill_match_score", 69),
                    "formatting_issues": data.get("formatting_issues", ["Long paragraphs", "Lack of bullet points"]),
                    "grammar_issues": data.get("grammar_issues", ["Missing commas", "Weak verbs"]),
                    "missing_hard_skills": data.get("missing_hard_skills", ["Docker", "Kubernetes"]),
                    "missing_soft_skills": data.get("missing_soft_skills", ["Leadership"]),
                    "recommendations": data.get("recommendations", [
                        "Add quantifiable achievements",
                        "Use stronger action verbs",
                        "Shorten sentences for readability"
                    ])
                }
            else:
                log.error(f"ATS API error: {response.status_code} - {response.text}")
        except Exception as e:
            log.error(f"ATS API request failed: {e}")
    # Fallback to mock data if API fails or no key
    return {
        "overall_score": 78,
        "keyword_score": 71,
        "skill_match_score": 69,
        "formatting_issues": ["Long paragraphs", "Lack of bullet points"],
        "grammar_issues": ["Missing commas", "Weak verbs"],
        "missing_hard_skills": ["Docker", "Kubernetes"],
        "missing_soft_skills": ["Leadership"],
        "recommendations": [
            "Add quantifiable achievements",
            "Use stronger action verbs",
            "Shorten sentences for readability"
        ]
    }


def extract_resume_structure(text: str):
    return {
        "name": "Your Name",
        "job_title": "Job Title",
        "summary": text[:300],
        "skills": ["Python", "React", "Node.js"],
        "experience": [
            {
                "position": "Software Developer",
                "company": "Tech Corp",
                "years": "2021-2024",
                "description": "Developed backend services and APIs."
            }
        ]
    }


@app.route("/enhance", methods=["POST"])
def enhance_resume():
    print("📥 /enhance called")

    file = request.files.get("resume_file")
    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    raw = file.read().decode("utf-8", errors="ignore")
    print(f"📄 Raw file length: {len(raw)}")

    try:
        enhanced = enhance_with_openai(raw)  # Using OpenAI for primary enhancement
        structured = extract_resume_structure(enhanced)
        ats = ats_analysis(raw)

        return jsonify({
            "original_text": raw,
            "enhanced_text": enhanced,
            "structured": structured,
            "ats": ats
        })

    except Exception as e:
        print("🔥 ERROR in /enhance:", e)
        return jsonify({"error": str(e)}), 500


@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    data = request.json
    template_name = data.get('template', 'modern_minimal')  # Default to modern_minimal
    template_path = os.path.join('backend', 'templates', f'{template_name}.tex')

    if not os.path.exists(template_path):
        return jsonify({"error": f"Template {template_name} not found"}), 404

    with open(template_path, 'r') as f:
        template_content = f.read()

    # Prepare data for template
    skills_list = '\n'.join([f'\\item {skill}' for skill in data.get("skills", [])])
    experience_list = '\n'.join([
        f'\\subsection{{{exp["position"]} at {exp["company"]} ({exp["years"]})}}\n{exp["description"]}'
        for exp in data.get("experience", [])
    ])

    template_data = {
        'name': data.get('name', 'Your Name'),
        'job_title': data.get('job_title', 'Job Title'),
        'summary': data.get('summary', ''),
        'skills': skills_list,
        'experience': experience_list
    }

    # Render template with Jinja2
    template = Template(template_content)
    rendered_tex = template.render(**template_data)

    # Compile to PDF
    with tempfile.TemporaryDirectory() as tmpdir:
        tex_file = os.path.join(tmpdir, 'resume.tex')
        pdf_file = os.path.join(tmpdir, 'resume.pdf')
        with open(tex_file, 'w') as f:
            f.write(rendered_tex)
        subprocess.run(['pdflatex', '-output-directory', tmpdir, tex_file], check=True)
        with open(pdf_file, 'rb') as f:
            pdf_data = f.read()

    return send_file(pdf_data, as_attachment=True, download_name='resume.pdf', mimetype='application/pdf')


@app.route("/export/docx", methods=["POST"])
def export_docx():
    try:
        from docx import Document
        data = request.json

        doc = Document()
        doc.add_heading(data.get("name"), level=0)
        doc.add_heading(data.get("job_title"), level=1)

        doc.add_heading("Summary", level=2)
        doc.add_paragraph(data.get("summary"))

        doc.add_heading("Skills", level=2)
        for s in data.get("skills", []):
            doc.add_paragraph(s, style="List Bullet")

        doc.add_heading("Experience", level=2)
        for exp in data.get("experience", []):
            doc.add_paragraph(
                f"{exp['position']} — {exp['company']} ({exp['years']})",
                style="List Bullet"
            )
            doc.add_paragraph(exp["description"])

        filename = "resume.docx"
        doc.save(filename)

        return send_file(filename, as_attachment=True)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/manual-entry", methods=["POST"])
def manual_entry():
    data = request.json
    # Process manual entry data
    text = f"{data.get('name')} - {data.get('job_title')}\n\nSummary: {data.get('summary')}\n\nSkills: {', '.join(data.get('skills', []))}\n\nExperience:\n"
    for exp in data.get('experience', []):
        text += f"{exp['position']} at {exp['company']} ({exp['years']}): {exp['description']}\n"
    enhanced = enhance_with_openai(text)
    ats = ats_analysis(text)
    return jsonify({"enhanced_text": enhanced, "ats": ats, "structured": data})

@app.route("/ats-score", methods=["POST"])
def ats_score():
    text = request.json.get("text")
    score = ats_analysis(text)
    return jsonify(score)

@app.route("/feedback-chat", methods=["POST"])
def feedback_chat():
    message = request.json.get("message")
    # Use OpenAI for chat feedback
    response = openai_client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": f"Provide feedback on this resume improvement suggestion: {message}"}],
        max_tokens=500
    )
    return jsonify({"response": response.choices[0].message.content.strip()})

@app.route("/score-tracker", methods=["GET"])
def score_tracker():
    # Placeholder for tracking score changes; in real app, store in DB
    return jsonify({"history": [{"date": "2023-10-01", "score": 75}, {"date": "2023-10-02", "score": 80}]})

@app.route("/download/<path:filename>")
def download(filename):
    if os.path.exists(filename):
        return send_file(filename, as_attachment=True)
    return jsonify({"error": "File not found"}), 404


if __name__ == "__main__":
    log.info("Starting Resume backend...")
    app.run(host="0.0.0.0", port=5000, debug=True)
